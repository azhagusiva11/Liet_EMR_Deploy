"""
Export tools for generating PDF and Word summaries
Handles both PDF and DOCX generation with proper formatting
"""

import os
import re
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus.flowables import HRFlowable
import json

# Import python-docx for Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Try to register a Unicode-supporting font
try:
    if os.name == 'nt':
        font_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/ArialUni.ttf",
            "C:/Windows/Fonts/NotoSans-Regular.ttf",
            "C:/Windows/Fonts/Calibri.ttf"
        ]
    else:
        font_paths = [
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
            "/System/Library/Fonts/Helvetica.ttc"
        ]
    
    font_registered = False
    for font_path in font_paths:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('UnicodeFont', font_path))
            font_registered = True
            break
    
    if not font_registered:
        print("Warning: No Unicode font found. Using default font.")
        
except Exception as e:
    print(f"Font registration warning: {e}")

def sanitize_text(text):
    """Sanitize text for PDF generation"""
    if text is None:
        return ""
    text = str(text)
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    return text

def create_header_footer(canvas, doc, patient_id=None, format_type=None):
    """Add header and footer to each page - NOW WITH FORMAT TYPE AND PATIENT ID"""
    canvas.saveState()
    
    # Header - Now includes format type
    canvas.setFont('Helvetica-Bold', 14)
    header_text = f"Smart EMR - Clinical Visit Summary ({format_type} Format)" if format_type else "Smart EMR - Clinical Visit Summary"
    canvas.drawCentredString(A4[0]/2, A4[1] - 50, header_text)
    
    # Footer - Now includes patient ID
    canvas.setFont('Helvetica', 9)
    canvas.drawString(inch, 0.75 * inch, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if patient_id:
        canvas.drawString(inch, 0.5 * inch, f"Patient ID: {patient_id}")
    canvas.drawRightString(A4[0] - inch, 0.75 * inch, f"Page {doc.page}")
    
    canvas.restoreState()

def analyze_vitals_criticality(vitals):
    """Analyze vitals and return critical findings"""
    critical_findings = []
    urgent_findings = []
    normal_findings = []
    
    if 'blood_pressure' in vitals and vitals['blood_pressure']:
        try:
            sys, dia = map(int, vitals['blood_pressure'].split('/'))
            if sys < 90 or dia < 60:
                critical_findings.append(f"🔴 CRITICAL Hypotension: {sys}/{dia} mmHg")
            elif sys < 100:
                urgent_findings.append(f"🟡 Hypotension: {sys}/{dia} mmHg")
            elif sys > 180 or dia > 120:
                critical_findings.append(f"🔴 CRITICAL Hypertension: {sys}/{dia} mmHg")
            elif sys > 140 or dia > 90:
                urgent_findings.append(f"🟡 Hypertension: {sys}/{dia} mmHg")
            else:
                normal_findings.append(f"✓ BP: {sys}/{dia} mmHg (Normal)")
        except:
            normal_findings.append(f"✓ BP: {vitals['blood_pressure']}")
    
    if 'heart_rate' in vitals and vitals['heart_rate']:
        hr = int(vitals['heart_rate'])
        if hr < 50 or hr > 120:
            critical_findings.append(f"🔴 CRITICAL Heart Rate: {hr} bpm")
        elif hr < 60 or hr > 100:
            urgent_findings.append(f"🟡 Heart Rate: {hr} bpm")
        else:
            normal_findings.append(f"✓ HR: {hr} bpm (Normal)")
    
    if 'temperature' in vitals and vitals['temperature']:
        temp = float(vitals['temperature'])
        if temp < 35.0 or temp > 40.0:
            critical_findings.append(f"🔴 CRITICAL Temperature: {temp}°C")
        elif temp < 36.0 or temp > 38.0:
            urgent_findings.append(f"🟡 Temperature: {temp}°C")
        else:
            normal_findings.append(f"✓ Temp: {temp}°C (Normal)")
    
    if 'spo2' in vitals and vitals['spo2']:
        spo2 = int(vitals['spo2'])
        if spo2 < 90:
            critical_findings.append(f"🔴 CRITICAL SpO2: {spo2}%")
        elif spo2 < 95:
            urgent_findings.append(f"🟡 SpO2: {spo2}%")
        else:
            normal_findings.append(f"✓ SpO2: {spo2}% (Normal)")
    
    if 'respiratory_rate' in vitals and vitals['respiratory_rate']:
        rr = int(vitals['respiratory_rate'])
        if rr < 10 or rr > 30:
            critical_findings.append(f"🔴 CRITICAL RR: {rr}/min")
        elif rr < 12 or rr > 20:
            urgent_findings.append(f"🟡 RR: {rr}/min")
        else:
            normal_findings.append(f"✓ RR: {rr}/min (Normal)")
    
    if 'weight' in vitals and vitals['weight']:
        normal_findings.append(f"✓ Weight: {vitals['weight']} kg")
    
    if 'height' in vitals and vitals['height']:
        normal_findings.append(f"✓ Height: {vitals['height']} cm")
        # Calculate BMI if both available
        if 'weight' in vitals and vitals['weight']:
            height_m = float(vitals['height']) / 100
            bmi = float(vitals['weight']) / (height_m * height_m)
            if bmi < 18.5:
                urgent_findings.append(f"🟡 Underweight: BMI {bmi:.1f}")
            elif bmi > 30:
                urgent_findings.append(f"🟡 Obese: BMI {bmi:.1f}")
            elif bmi > 25:
                urgent_findings.append(f"🟡 Overweight: BMI {bmi:.1f}")
            else:
                normal_findings.append(f"✓ BMI: {bmi:.1f} (Normal)")
    
    return critical_findings, urgent_findings, normal_findings

def format_prescription_properly(prescription_text):
    """Format prescription in a structured way - FIXED VERSION"""
    if not prescription_text or prescription_text.strip() == "":
        return []
    
    # Clean the text first
    clean_text = prescription_text.strip()
    
    formatted_meds = []
    lines = clean_text.split('\n')
    
    med_counter = 1
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip non-medication lines that might have leaked through
        skip_phrases = ['teaching point', 'safety advice', 'follow-up', 'appointment', 
                       'return immediately', 'avoid', 'crucial', 'progression']
        if any(phrase in line.lower() for phrase in skip_phrases):
            continue
            
        # Remove any leading dash or weird formatting
        if line.startswith('- '):
            line = line[2:]
        
        # Fix numbering - always renumber
        if line and line[0].isdigit() and len(line) > 1 and line[1] == '.':
            line = line[2:].strip()  # Remove old number
        
        # Skip empty lines after cleaning
        if not line:
            continue
            
        # Add proper numbering
        formatted_line = f"{med_counter}. {line}"
        med_counter += 1
        
        formatted_meds.append({
            'medication': line,
            'formatted': formatted_line
        })
    
    return formatted_meds

def generate_visit_pdf(patient_data, visit_data):
    """Generate a PDF for a specific visit - NOW WITH FORMAT TYPE AND PATIENT ID"""
    visit_id = visit_data.get('visit_id', datetime.now().strftime('%Y%m%d%H%M%S'))
    patient_name = patient_data.get('name', 'patient').replace(' ', '_')
    patient_id = patient_data.get('id', 'Unknown')
    format_type = visit_data.get('format_type', 'SOAP')  # Get format type from visit data
    
    filename = f"visit_summary_{patient_name}_{visit_id}.pdf"
    
    filepath = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)
    
    # Create custom header/footer function with patient ID and format
    def custom_header_footer(canvas, doc):
        create_header_footer(canvas, doc, patient_id, format_type)
    
    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=100,
        bottomMargin=80  # Increased for patient ID
    )
    
    # Define styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f4e79'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#2c5282'),
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        wordWrap='CJK'  # Better word wrapping
    )
    
    # Build content
    content = []
    
    # Title - Now includes format type
    title_text = f"Clinical Visit Summary - {format_type} Format"
    content.append(Paragraph(title_text, title_style))
    content.append(Spacer(1, 0.2*inch))
    
    # Patient Information - Now includes ID
    content.append(Paragraph("Patient Information", heading_style))
    
    patient_info = [
        ["Patient ID:", sanitize_text(patient_id)],
        ["Name:", sanitize_text(patient_data.get('name', 'N/A'))],
        ["Age:", f"{patient_data.get('age', 'N/A')} years"],
        ["Sex:", sanitize_text(patient_data.get('sex', 'N/A'))],
        ["Mobile:", sanitize_text(patient_data.get('mobile', 'N/A'))],
        ["Visit Date:", visit_data.get('timestamp', 'N/A')[:10]],
    ]
    
    if patient_data.get('blood_group'):
        patient_info.append(["Blood Group:", sanitize_text(patient_data.get('blood_group'))])
    
    if patient_data.get('allergies'):
        allergies_text = ", ".join(patient_data.get('allergies', []))
        patient_info.append(["⚠️ ALLERGIES:", sanitize_text(allergies_text)])
    
    patient_table = Table(patient_info, colWidths=[2*inch, 4*inch])
    patient_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#2c5282')),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    
    content.append(patient_table)
    content.append(Spacer(1, 0.3*inch))
    
    # Vital Signs
    if 'vitals' in visit_data and visit_data['vitals']:
        content.append(Paragraph("Vital Signs", heading_style))
        
        critical_findings, urgent_findings, normal_findings = analyze_vitals_criticality(visit_data['vitals'])
        
        # Show all findings
        all_findings = critical_findings + urgent_findings + normal_findings
        for finding in all_findings:
            content.append(Paragraph(finding, normal_style))
        
        content.append(Spacer(1, 0.3*inch))
    
    # Clinical Summary (WITHOUT prescription)
    if 'summary' in visit_data and visit_data['summary']:
        content.append(Paragraph("Clinical Assessment", heading_style))
        
        summary_text = sanitize_text(visit_data['summary'])
        # Remove any prescription that might have leaked into summary
        for marker in ['PRESCRIPTION:', 'Prescription:', 'TREATMENT:', '===PRESCRIPTION', '=== PRESCRIPTION START ===']:
            if marker in summary_text:
                summary_text = summary_text.split(marker)[0].strip()
        
        # Create paragraphs for better formatting
        summary_paragraphs = summary_text.split('\n\n')
        for para in summary_paragraphs:
            if para.strip():
                content.append(Paragraph(para.strip(), normal_style))
                content.append(Spacer(1, 0.1*inch))
        
        content.append(Spacer(1, 0.2*inch))
    
    # Prescription Section (SEPARATE) - FIXED
    if 'prescription' in visit_data and visit_data['prescription'] and visit_data['prescription'].strip():
        content.append(Paragraph("Prescription", heading_style))
        
        # Debug print
        print(f"DEBUG: Raw prescription text: {visit_data['prescription']}")
        
        formatted_meds = format_prescription_properly(visit_data['prescription'])
        
        if formatted_meds:
            rx_data = []
            for med in formatted_meds:
                parts = med['formatted'].split('.', 1)
                if len(parts) == 2:
                    rx_data.append([parts[0] + ".", parts[1].strip()])
                else:
                    rx_data.append(["", med['formatted']])
            
            # FIXED: Wider column + word wrap
            rx_table = Table(rx_data, colWidths=[0.4*inch, 5.8*inch])
            rx_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('PADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f5f5f5')),
                ('BOX', (0, 0), (-1, -1), 1, colors.grey),
                ('WORDWRAP', (1, 0), (1, -1), 'CJK'),  # Better word wrap
            ]))
            content.append(rx_table)
        else:
            # If no formatted meds but prescription exists, show raw text
            content.append(Paragraph("Prescription details:", normal_style))
            prescription_para = Paragraph(sanitize_text(visit_data['prescription']), normal_style)
            content.append(prescription_para)
    else:
        # No prescription section at all
        content.append(Paragraph("Prescription", heading_style))
        content.append(Paragraph("No prescription provided", normal_style))
        
    content.append(Spacer(1, 0.3*inch))
    
    # Doctor's signature
    content.append(Spacer(1, 0.5*inch))
    content.append(HRFlowable(width=2*inch, thickness=1, color=colors.black))
    content.append(Paragraph("Doctor's Signature & Stamp", normal_style))
    
    # Build PDF
    try:
        doc.build(content, onFirstPage=custom_header_footer, onLaterPages=custom_header_footer)
        print(f"PDF generated successfully: {filepath}")
        return filepath
    except Exception as e:
        print(f"Error generating PDF: {e}")
        # Fallback to text file
        text_path = filepath.replace('.pdf', '.txt')
        try:
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(f"CLINICAL VISIT SUMMARY - {format_type} FORMAT\n")
                f.write("="*50 + "\n\n")
                f.write(f"Patient ID: {patient_id}\n")
                f.write(f"Patient: {patient_data.get('name', 'N/A')}\n")
                f.write(f"Age: {patient_data.get('age', 'N/A')} years\n")
                f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
                
                if visit_data.get('summary'):
                    f.write("CLINICAL ASSESSMENT:\n")
                    f.write(visit_data['summary'] + "\n\n")
                
                if visit_data.get('prescription'):
                    f.write("PRESCRIPTION:\n")
                    f.write(visit_data['prescription'] + "\n")
            
            print(f"Saved as text file: {text_path}")
            return text_path
        except Exception as e2:
            print(f"Failed to save as text: {e2}")
            return None


def generate_visit_docx(patient_data, visit_data):
    """Generate a Word document for a specific visit"""
    visit_id = visit_data.get('visit_id', datetime.now().strftime('%Y%m%d%H%M%S'))
    patient_name = patient_data.get('name', 'patient').replace(' ', '_')
    patient_id = patient_data.get('id', 'Unknown')
    format_type = visit_data.get('format_type', 'SOAP')
    
    filename = f"visit_summary_{patient_name}_{visit_id}.docx"
    filepath = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)
    
    # Create document
    doc = Document()
    
    # Set up styles
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(31, 78, 121)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(18)
    
    heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Calibri'
    heading_font.size = Pt(16)
    heading_font.bold = True
    heading_font.color.rgb = RGBColor(44, 82, 130)
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"Smart EMR - Clinical Visit Summary ({format_type} Format)"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title
    title = doc.add_paragraph(f"Clinical Visit Summary - {format_type} Format", style='CustomTitle')
    
    # Patient Information
    doc.add_paragraph("Patient Information", style='CustomHeading')
    
    # Add patient details table
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # Add rows
    def add_info_row(label, value):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    add_info_row("Patient ID:", patient_id)
    add_info_row("Name:", patient_data.get('name', 'N/A'))
    add_info_row("Age:", f"{patient_data.get('age', 'N/A')} years")
    add_info_row("Sex:", patient_data.get('sex', 'N/A'))
    add_info_row("Mobile:", patient_data.get('mobile', 'N/A'))
    add_info_row("Visit Date:", visit_data.get('timestamp', 'N/A')[:10])
    
    if patient_data.get('blood_group'):
        add_info_row("Blood Group:", patient_data.get('blood_group'))
    
    if patient_data.get('allergies'):
        add_info_row("⚠️ ALLERGIES:", ", ".join(patient_data.get('allergies', [])))
    
    doc.add_paragraph()  # Add spacing
    
    # Vital Signs
    if 'vitals' in visit_data and visit_data['vitals']:
        doc.add_paragraph("Vital Signs", style='CustomHeading')
        
        critical_findings, urgent_findings, normal_findings = analyze_vitals_criticality(visit_data['vitals'])
        
        # Add all findings
        all_findings = critical_findings + urgent_findings + normal_findings
        for finding in all_findings:
            para = doc.add_paragraph(finding)
            if "🔴" in finding:
                para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            elif "🟡" in finding:
                para.runs[0].font.color.rgb = RGBColor(255, 165, 0)
    
    # Clinical Summary
    if 'summary' in visit_data and visit_data['summary']:
        doc.add_paragraph("Clinical Assessment", style='CustomHeading')
        
        summary_text = visit_data['summary']
        # Remove any prescription that might have leaked into summary
        for marker in ['PRESCRIPTION:', 'Prescription:', 'TREATMENT:', '===PRESCRIPTION']:
            if marker in summary_text:
                summary_text = summary_text.split(marker)[0].strip()
        
        # Add summary paragraphs
        for para in summary_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Prescription
    if 'prescription' in visit_data and visit_data['prescription']:
        doc.add_paragraph("Prescription", style='CustomHeading')
        
        formatted_meds = format_prescription_properly(visit_data['prescription'])
        
        if formatted_meds:
            for med in formatted_meds:
                para = doc.add_paragraph(med['formatted'])
                para.paragraph_format.left_indent = Inches(0.25)
        else:
            doc.add_paragraph(visit_data['prescription'])
    else:
        doc.add_paragraph("Prescription", style='CustomHeading')
        doc.add_paragraph("No prescription provided")
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("_" * 50)
    doc.add_paragraph("Doctor's Signature & Stamp")
    
    # Add footer with generation time
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Patient ID: {patient_id}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    try:
        doc.save(filepath)
        print(f"Word document generated successfully: {filepath}")
        return filepath
    except Exception as e:
        print(f"Error generating Word document: {e}")
        return None


def generate_discharge_summary(patient_data, admission_data, discharge_data):
    """Generate a discharge summary PDF"""
    patient_name = patient_data.get('name', 'patient').replace(' ', '_')
    filename = f"discharge_summary_{patient_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
    
    combined_data = {
        'timestamp': discharge_data.get('discharge_date', datetime.now().strftime('%Y-%m-%d')),
        'visit_type': 'Discharge Summary',
        'chief_complaint': admission_data.get('chief_complaint', 'N/A'),
        'summary': f"ADMISSION DETAILS:\n{admission_data.get('summary', '')}\n\nDISCHARGE SUMMARY:\n{discharge_data.get('summary', '')}",
        'prescription': discharge_data.get('discharge_prescription', ''),
        'vitals': discharge_data.get('vitals_at_discharge', {}),
        'format_type': 'Discharge'  # Special format for discharge
    }
    
    return generate_visit_pdf(patient_data, combined_data)